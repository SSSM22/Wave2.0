from django.shortcuts import render,redirect
import requests
from bs4 import BeautifulSoup
from django.http import HttpResponse
import json
import gc
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
import time
from selenium.webdriver.chrome.options import Options
from res.image import scrape_website
from docx import Document
import shutil
import mimetypes
from django.http import FileResponse, Http404
import os
from django.views.decorators.csrf import csrf_exempt
from .models import Users

user ={}

# Create your views here.
service = ChromeService(ChromeDriverManager().install())
def index(request):
    return render(request,'index.html')

@csrf_exempt
def details(request):
 if request.method == 'POST':
    headers = {"cookie": "CONSENT=YES+cb.20230531-04-p0.en+FX+908"}
    body_unicode = request.body.decode('utf-8')
    body = json.loads(body_unicode)
    link = body['link']
    # print(link)
    # print(body)
    user = body['user']
    # gc.disable()
    # link='https://internshala.com/internship/detail/work-from-home-php-development-internship-at-areeyarath-boonsoontornpat1718254017'
    if 'internshala' in link:
        response=requests.get(link,headers=headers)
        print(response.status_code)
        if response.status_code == 200:
            # print(user)
            soup = BeautifulSoup(response.content, 'html.parser')
            about = soup.find('div',class_="internship_details")
            todo = about.find_all('div',class_="text-container")
            skills=about.find_all('div',class_="round_tabs_container")
            first_todo = todo[0] if todo else None
            first_skills = skills[0] if skills else None
            
            if first_todo:
                print("First div in todo:")
                first_todo=first_todo.get_text()
            if first_skills:
            
                first_skills=first_skills.get_text()
                print("First div in skills:",first_skills)
            text=generate_content(first_todo,first_skills)
            print(user)
            shutil.copy(r'res\sesume_template.docx', r'res\download.docx')
            f=user['name']
            m=user['email']
            p=user['contact']
            link=user['linkedin']
            git=user['github']
            clgname=user['education'].split(',')[0]

            graduation=user['education'].split(',')[1]
            second=user['education'].split(',')[2]
            s_year=user['education'].split(',')[3]
            tp1=user['project1'].split(',')[0]
            tp2=user['project2'].split(',')[0]
            tp3=user['project3'].split(',')[0]
            dp1=user['project1'].split(',')[1]
            dp2=user['project2'].split(',')[1]
            dp3=user['project3'].split(',')[1]
            ach1=user['achievements'].split(',')[0]
            ach2=user['achievements'].split(',')[1]
            ach3=user['achievements'].split(',')[2]
            skill=user['skills'].split(';')[0]
            wf=user['skills'].split(';')[1]
            company=user['experience'].split(',')[1]
            exp_t=user['experience'].split(',')[0]
            year=user['experience'].split(',')[2]
            place=user['place']
            # Fill the template
            data = json.loads(text)
            
            text = data['candidates'][0]['content']['parts'][0]['text']
            prefix_to_remove = "**Objective:**"
            text = text.replace(prefix_to_remove, "")
            print(text)
            document = Document('res\download.docx')
            for paragraph in document.paragraphs:
                if 'wf' in paragraph.text:
                    paragraph.text = paragraph.text.replace('wf',wf)
                if 'second' in paragraph.text:
                    paragraph.text = paragraph.text.replace('second',second)
                if 's_y' in paragraph.text:
                    paragraph.text = paragraph.text.replace('s_y',s_year)
                if 'place_of_job' in paragraph.text:
                    paragraph.text = paragraph.text.replace('place_of_job',place)
                if 'Object.ive' in paragraph.text:
                    paragraph.text=paragraph.text.replace('Object.ive',text)
                if 'name' in paragraph.text:
                    paragraph.text = paragraph.text.replace('name', f)
                if 'contact' in paragraph.text:
                    paragraph.text = paragraph.text.replace('contact', m)
                if 'linkedin' in paragraph.text:
                    paragraph.text = paragraph.text.replace('linkedin',link)
                if 'github' in paragraph.text:
                    paragraph.text = paragraph.text.replace('github',git)
                if 'phone_number' in paragraph.text:
                    paragraph.text = paragraph.text.replace('phone_number', str(p))
                if 'batchlor_degree' in paragraph.text:
                    paragraph.text = paragraph.text.replace('batchlor_degree', clgname)
                if 'graduation' in paragraph.text:
                    paragraph.text = paragraph.text.replace('graduation',graduation)
                if 'lang' in paragraph.text:
                    paragraph.text = paragraph.text.replace('lang',skill)
                if 'Expe_t' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Expe_t', exp_t)
                if 'year' in paragraph.text:
                    paragraph.text = paragraph.text.replace('year',year)
                if 'cmp' in paragraph.text:
                    paragraph.text = paragraph.text.replace('cmp',company)
                if 'Project_1t' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Project_1t',tp1)
                if 'Project_2t' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Project_2t',tp2)
                if 'Project_3t' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Project_3t',tp3)
                if 'Project_1d' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Project_1d',dp1)
                if 'Project_2d' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Project_2d',dp2)
                if 'Project_3d' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Project_3d',dp3)
                if 'Achieve1' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Achieve1',ach1)
                if 'Achieve2' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Achieve2',ach2)
                if 'Achieve3' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Achieve3',ach3)    
            # Save the filled document
            document.save('res\download.docx')
            file_path = os.path.join('res/','download.docx')
            if os.path.exists(file_path):
                response = FileResponse(open(file_path, 'rb'))
                return response
            return resume

    if 'naukri' in link:
        ans=scrape_website(link)
        text=generate_content(ans)
        resume=update_resume(text)
        return resume

        

    else:
            response=requests.get(link,headers=headers)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                generate_content(soup)

            return HttpResponse("yes")

def generate_content(job_des,skills=None):
    gemini_api_key = 'AIzaSyBzsemF33e6V-x2vz3pOtSKWwsvQ1h52v4'#'AIzaSyCqdMnGnigseAJH8bBSfuggxF-oqU5lrXQ'  # Replace with your actual API key
    url = f'https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key={gemini_api_key}'

    if skills!=None:

        data = {
            "contents": [
                {
                    "parts": [
                        {
                         "text": f'make an objective in 30 words for my resume with job description {job_des} and skills {skills}'
                        }
                    ]
                }
            ]
        }

        headers = {
            'Content-Type': 'application/json'
        }

        try:
            response = requests.post(url, json=data, headers=headers)
            response.raise_for_status()  # Raise an error for bad status codes

            return(response.text)

        except requests.exceptions.RequestException as e:
            print("error")
            # return {'error': f'Request error: {str(e)}'}
    else:
        data = {
            "contents": [
                {
                    "parts": [
                        {
                         "text": f'make my resume with job description avaliable in html page along with skills required {job_des}'
                        }
                    ]
                }
            ]
        }

        headers = {
            'Content-Type': 'application/json'
        }

        try:
            response = requests.post(url, json=data, headers=headers)
            response.raise_for_status()  # Raise an error for bad status codes

            return(response.text)

        except requests.exceptions.RequestException as e:
            print("error")
            # return {'error': f'Request error: {str(e)}'}





# def update_resume(text,user):
            
            
            # shutil.copy(r'res\sesume_template.docx', r'res\download.docx')
            # f=user['name']
            # m=user['email']
            # p=user['contact']
            # link=user['linkedin']
            # git=user['github']
            # clgname=user['education'].split(',')[0]

            # graduation=user['education'].split(',')[1]
            # second=user['education'].split(',')[2]
            # s_year=user['education'].split(',')[3]
            # tp1=user['project1'].split(',')[0]
            # tp2=user['project2'].split(',')[0]
            # tp3=user['project3'].split(',')[0]
            # dp1=user['project1'].split(',')[1]
            # dp2=user['project2'].split(',')[1]
            # dp3=user['project3'].split(',')[1]
            # ach1=user['achievements'].split(',')[0]
            # ach2=user['achievements'].split(',')[1]
            # ach3=user['achievements'].split(',')[2]
            # skill=user['skills'].split(';')[0]
            # wf=user['skills'].split(';')[1]
            # company=user['experience'].split(',')[1]
            # exp_t=user['experience'].split(',')[0]
            # year=user['experience'].split(',')[2]
            # place=user['place']
            # # Fill the template
            # data = json.loads(text)
            
            # text = data['candidates'][0]['content']['parts'][0]['text']
            # prefix_to_remove = "**Objective:**"
            # text = text.replace(prefix_to_remove, "")
            # print(text)
            # document = Document('res\download.docx')
            # for paragraph in document.paragraphs:
            #     if 'wf' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('wf',wf)
            #     if 'second' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('second',second)
            #     if 's_y' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('s_y',s_year)
            #     if 'place_of_job' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('place_of_job',place)
            #     if 'Object.ive' in paragraph.text:
            #         paragraph.text=paragraph.text.replace('Object.ive',text)
            #     if 'name' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('name', f)
            #     if 'contact' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('contact', m)
            #     if 'linkedin' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('linkedin',link)
            #     if 'github' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('github',git)
            #     if 'phone_number' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('phone_number', str(p))
            #     if 'batchlor_degree' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('batchlor_degree', clgname)
            #     if 'graduation' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('graduation',graduation)
            #     if 'lang' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('lang',skill)
            #     if 'Expe_t' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Expe_t', exp_t)
            #     if 'year' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('year',year)
            #     if 'cmp' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('cmp',company)
            #     if 'Project_1t' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Project_1t',tp1)
            #     if 'Project_2t' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Project_2t',tp2)
            #     if 'Project_3t' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Project_3t',tp3)
            #     if 'Project_1d' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Project_1d',dp1)
            #     if 'Project_2d' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Project_2d',dp2)
            #     if 'Project_3d' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Project_3d',dp3)
            #     if 'Achieve1' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Achieve1',ach1)
            #     if 'Achieve2' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Achieve2',ach2)
            #     if 'Achieve3' in paragraph.text:
            #         paragraph.text = paragraph.text.replace('Achieve3',ach3)    
            # # Save the filled document
            # document.save('res\download.docx')
            # file_path = os.path.join('res/','download.docx')
            # if os.path.exists(file_path):
            #     response = FileResponse(open(file_path, 'rb'))
            #     return response
            

